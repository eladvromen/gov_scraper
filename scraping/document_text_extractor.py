import requests
import pandas as pd
import json
import os
import time
import logging
from datetime import datetime, timedelta
from pathlib import Path
from typing import Dict, List, Optional, Tuple
import tempfile
import random
from urllib.parse import urlparse
from bs4 import BeautifulSoup
import re
import io
from tqdm import tqdm

# Document processing libraries
try:
    import docx2txt
    DOCX2TXT_AVAILABLE = True
except ImportError:
    DOCX2TXT_AVAILABLE = False

try:
    from docx import Document
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False

try:
    import PyPDF2
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False

try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False

class DocumentTextExtractor:
    """
    Extracts text from downloadable Word/PDF documents for UK Immigration Tribunal cases
    where decision_text is missing from the original scrape.
    """
    
    def __init__(self, 
                 csv_path: str,
                 json_dir: str = "data/json",
                 rate_limit: float = 0.5,  # Start with faster rate limit
                 max_retries: int = 3,
                 test_mode: bool = False,
                 test_count: int = 5,
                 batch_size: int = 500,
                 checkpoint_dir: str = "checkpoints"):
        """
        Initialize the document text extractor.
        
        Args:
            csv_path: Path to the processed CSV file
            json_dir: Directory containing JSON files to update
            rate_limit: Time in seconds to wait between downloads
            max_retries: Maximum retry attempts for failed downloads
            test_mode: If True, only process a small number of files for testing
            test_count: Number of files to process in test mode
            batch_size: Number of cases to process per batch (for large runs)
            checkpoint_dir: Directory to store progress checkpoints
        """
        self.csv_path = Path(csv_path)
        self.json_dir = Path(json_dir)
        self.rate_limit = rate_limit
        self.initial_rate_limit = rate_limit  # Store original for adaptive adjustment
        self.max_retries = max_retries
        self.test_mode = test_mode
        self.test_count = test_count
        self.batch_size = batch_size
        self.checkpoint_dir = Path(checkpoint_dir)
        
        # Create checkpoint directory
        self.checkpoint_dir.mkdir(exist_ok=True)
        
        # Progress tracking
        self.processed_cases = set()
        self.successful_cases = []
        self.failed_cases = []
        self.current_batch = 0
        
        # Performance tracking
        self.start_time = None
        self.cases_processed = 0
        self.total_download_time = 0
        self.total_extraction_time = 0
        
        # Create session for requests
        self.session = requests.Session()
        self.session.headers.update({
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        })
        
        # Create temp directory for downloads
        self.temp_dir = Path(tempfile.mkdtemp(prefix="tribunal_docs_"))
        
        # Setup logging
        self.setup_logging()
        
        # Check available libraries
        self.check_dependencies()
        
        # Load CSV data
        self.load_missing_cases()
        
    def setup_logging(self):
        """Set up logging configuration with minimal console output for better progress bar visibility"""
        self.logger = logging.getLogger("document_extractor")
        self.logger.setLevel(logging.INFO)
        
        # Prevent duplicate handlers
        if not self.logger.handlers:
            # Create logs directory if it doesn't exist
            os.makedirs("logs", exist_ok=True)
            
            # File handler with UTF-8 encoding - DETAILED logging
            file_handler = logging.FileHandler(
                f"logs/document_extractor_{datetime.now().strftime('%Y%m%d_%H%M%S')}.log",
                encoding='utf-8'
            )
            file_handler.setLevel(logging.INFO)  # Detailed file logging
            
            # Console handler - MINIMAL logging for progress bar visibility
            console_handler = logging.StreamHandler()
            console_handler.setLevel(logging.WARNING)  # Only warnings and errors on console
            
            # Detailed formatter for file
            file_formatter = logging.Formatter('%(asctime)s - %(name)s - %(levelname)s - %(message)s')
            file_handler.setFormatter(file_formatter)
            
            # Minimal formatter for console
            console_formatter = logging.Formatter('%(levelname)s: %(message)s')
            console_handler.setFormatter(console_formatter)
            
            self.logger.addHandler(file_handler)
            self.logger.addHandler(console_handler)
            
            # Store handler references for dynamic adjustment
            self.file_handler = file_handler
            self.console_handler = console_handler
    
    def set_console_verbosity(self, verbose: bool = False):
        """
        Adjust console logging verbosity
        
        Args:
            verbose: If True, show INFO messages on console. If False, only WARNING/ERROR
        """
        if hasattr(self, 'console_handler'):
            if verbose:
                self.console_handler.setLevel(logging.INFO)
                print("🔊 Console logging: VERBOSE (INFO level)")
            else:
                self.console_handler.setLevel(logging.WARNING)
                print("🔇 Console logging: QUIET (WARNING level only)")
                print("📄 Detailed logs available in: logs/document_extractor_*.log")
    
    def check_dependencies(self):
        """Check which document processing libraries are available"""
        missing_libs = []
        
        if not DOCX2TXT_AVAILABLE and not PYTHON_DOCX_AVAILABLE:
            missing_libs.append("docx2txt or python-docx (for Word documents)")
        
        if not PYPDF2_AVAILABLE and not PDFPLUMBER_AVAILABLE:
            missing_libs.append("PyPDF2 or pdfplumber (for PDF documents)")
        
        if missing_libs:
            self.logger.warning(f"Missing libraries: {', '.join(missing_libs)}")
            self.logger.warning("Install with: pip install docx2txt python-docx PyPDF2 pdfplumber")
        
        self.logger.info(f"Available libraries: "
                        f"docx2txt={DOCX2TXT_AVAILABLE}, "
                        f"python-docx={PYTHON_DOCX_AVAILABLE}, "
                        f"PyPDF2={PYPDF2_AVAILABLE}, "
                        f"pdfplumber={PDFPLUMBER_AVAILABLE}")
    
    def load_missing_cases(self):
        """Load cases with missing decision text from CSV"""
        if not self.csv_path.exists():
            raise FileNotFoundError(f"CSV file not found: {self.csv_path}")
        
        self.logger.info(f"Loading CSV from {self.csv_path}")
        df = pd.read_csv(self.csv_path)
        
        # Find cases with missing decision text
        missing_text = (
            df['decision_text_cleaned'].isna() | 
            (df['decision_text_cleaned'] == '') |
            (df['decision_text_cleaned'].str.len() < 50)  # Very short text likely means missing
        )
        
        self.missing_cases = df[missing_text].copy()
        
        # Filter to cases that have the main URL (we'll get fresh download links from there)
        has_main_url = self.missing_cases['url'].notna()
        
        self.missing_cases = self.missing_cases[has_main_url].copy()
        
        if self.test_mode:
            self.missing_cases = self.missing_cases.head(self.test_count)
            self.logger.info(f"TEST MODE: Processing only {len(self.missing_cases)} cases")
        
        self.logger.info(f"Found {len(self.missing_cases)} cases with missing text and main URLs")
        
        if len(self.missing_cases) == 0:
            self.logger.warning("No cases found with missing text and main URLs!")
    
    def get_fresh_download_urls(self, case_url: str) -> Tuple[Optional[str], Optional[str]]:
        """
        Get fresh download URLs from the case page
        
        Args:
            case_url: The main case URL
            
        Returns:
            Tuple of (word_url, pdf_url) or (None, None) if failed
        """
        try:
            self.logger.info(f"Fetching fresh download URLs from {case_url}")
            
            # Add jitter to rate limit
            jitter = random.uniform(-0.2, 0.2) * self.rate_limit
            actual_delay = max(0.1, self.rate_limit + jitter)
            time.sleep(actual_delay)
            
            response = self.session.get(case_url, timeout=30)
            response.raise_for_status()
            
            soup = BeautifulSoup(response.text, 'html.parser')
            
            # Look for download links
            word_url = None
            pdf_url = None
            
            # Debug: log the HTML structure around download links
            download_div = soup.find('div', class_='download-links')
            if download_div:
                self.logger.info(f"Found download-links div with {len(download_div.find_all('a'))} links")
                # Log all links in the download div for debugging
                for link in download_div.find_all('a', href=True):
                    href = link.get('href', '')
                    text = link.get_text().strip()
                    self.logger.info(f"Download link: '{text}' -> '{href}'")
            
            # Find actual download links - look for AWS S3 or direct file links
            all_links = soup.find_all('a', href=True)
            for link in all_links:
                href = link.get('href', '')
                text = link.get_text().lower().strip()
                
                # Skip if this is the case page URL itself
                if case_url in href:
                    continue
                
                # Look for Word document patterns
                if (('word' in text and 'download' in text) or 
                    ('doc' in href.lower() and 's3' in href.lower()) or
                    ('.doc' in href.lower() and 'amazonaws' in href.lower())):
                    word_url = href
                    if not word_url.startswith(('http://', 'https://')):
                        word_url = f"https://tribunalsdecisions.service.gov.uk{word_url}"
                    self.logger.info(f"Found Word URL: {word_url}")
                    break
            
            # Find PDF link - look for actual PDF download URLs
            for link in all_links:
                href = link.get('href', '')
                text = link.get_text().lower().strip()
                
                # Skip if this is the case page URL itself
                if case_url in href:
                    continue
                
                # Look for PDF patterns - must be actual file downloads
                if (('pdf' in text and 'download' in text) or
                    ('.pdf' in href.lower() and 's3' in href.lower()) or
                    ('.pdf' in href.lower() and 'amazonaws' in href.lower())):
                    pdf_url = href
                    if not pdf_url.startswith(('http://', 'https://')):
                        pdf_url = f"https://tribunalsdecisions.service.gov.uk{pdf_url}"
                    self.logger.info(f"Found PDF URL: {pdf_url}")
                    break
            
            # If we didn't find specific PDF link but found Word link that's actually a PDF, that's ok
            if not pdf_url and word_url and '.pdf' in word_url.lower():
                self.logger.info("Word URL appears to be a PDF file")
            
            self.logger.info(f"Final URLs - Word: {'Yes' if word_url else 'No'}, PDF: {'Yes' if pdf_url else 'No'}")
            return word_url, pdf_url
            
        except Exception as e:
            self.logger.error(f"Error getting fresh URLs from {case_url}: {str(e)}")
            return None, None
    
    def download_file(self, url: str, filename: str) -> Optional[Path]:
        """
        Download a file from URL to temp directory
        
        Args:
            url: Download URL
            filename: Local filename to save as
            
        Returns:
            Path to downloaded file or None if failed
        """
        if not url or pd.isna(url):
            return None
        
        # Add jitter to rate limit
        jitter = random.uniform(-0.2, 0.2) * self.rate_limit
        actual_delay = max(0.1, self.rate_limit + jitter)
        time.sleep(actual_delay)
        
        file_path = self.temp_dir / filename
        
        for attempt in range(self.max_retries + 1):
            try:
                self.logger.info(f"Downloading {url} (attempt {attempt + 1})")
                
                response = self.session.get(url, timeout=60, stream=True)
                
                # Handle rate limiting
                if response.status_code == 429:
                    wait_time = int(response.headers.get('Retry-After', 60))
                    self.logger.warning(f"Rate limited. Waiting {wait_time} seconds.")
                    time.sleep(wait_time)
                    continue
                
                response.raise_for_status()
                
                # Check if we got HTML instead of a document (redirect to login page, etc.)
                content_type = response.headers.get('content-type', '').lower()
                if 'text/html' in content_type:
                    self.logger.warning(f"Got HTML response instead of document for {url}")
                    continue
                
                # Save file
                with open(file_path, 'wb') as f:
                    for chunk in response.iter_content(chunk_size=8192):
                        if chunk:
                            f.write(chunk)
                
                file_size = file_path.stat().st_size
                self.logger.info(f"Downloaded {filename} ({file_size:,} bytes)")
                
                # Basic file validation
                if file_size < 100:  # Suspiciously small file
                    self.logger.warning(f"Downloaded file {filename} is very small ({file_size} bytes)")
                    return None
                
                return file_path
                
            except Exception as e:
                if attempt < self.max_retries:
                    wait_time = (2 ** attempt) + random.uniform(0, 1)
                    self.logger.warning(f"Download failed ({str(e)}). Retrying in {wait_time:.2f} seconds...")
                    time.sleep(wait_time)
                else:
                    self.logger.error(f"Failed to download {url} after {self.max_retries} retries: {str(e)}")
                    return None
    
    def extract_text_from_document(self, file_path: str, file_url: str) -> Optional[str]:
        """
        Extract text from a document file
        
        Args:
            file_path: Local path to the downloaded file
            file_url: Original URL of the file
            
        Returns:
            Extracted text or None if extraction failed
        """
        try:
            # Check file size
            file_size = os.path.getsize(file_path)
            self.logger.info(f"Processing file: {file_size:,} bytes")
            
            if file_size == 0:
                self.logger.error("File is empty")
                return None
            
            if file_size > 50 * 1024 * 1024:  # 50MB limit
                self.logger.error(f"File too large: {file_size:,} bytes")
                return None
            
            # Read first few bytes to determine file type
            with open(file_path, 'rb') as f:
                header = f.read(512)
            
            # Detect file format from header
            if header.startswith(b'%PDF'):
                self.logger.info("Detected PDF file")
                return self._extract_from_pdf(file_path)
            elif header.startswith(b'PK\x03\x04'):
                # This is a ZIP-based file - check if it's a .docx
                if (b'word/' in header or 
                    b'[Content_Types].xml' in header or 
                    file_path.lower().endswith('.docx')):
                    self.logger.info("Detected modern Word document (.docx)")
                    return self._extract_from_docx(file_path)
                else:
                    self.logger.warning(f"Unknown ZIP-based file format. Header: {header[:100]}")
                    return self._extract_as_text(file_path)
            elif b'\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1' in header[:8]:
                self.logger.info("Detected old Word document (.doc) - OLE compound")
                return self._extract_from_old_doc(file_path)
            elif b'%PDF' in header:
                self.logger.info("Detected embedded PDF in document")
                return self._extract_embedded_pdf(file_path)
            else:
                self.logger.warning(f"Unknown file format, attempting text extraction. Header: {header[:50]}")
                return self._extract_as_text(file_path)
            
        except Exception as e:
            self.logger.error(f"Error extracting text from {file_path}: {str(e)}")
            return None
    
    def _extract_from_pdf(self, file_path: str) -> Optional[str]:
        """Extract text from PDF file"""
        # Try pdfplumber first (best for complex layouts)
        if PDFPLUMBER_AVAILABLE:
            try:
                import pdfplumber
                self.logger.info("Trying pdfplumber for PDF extraction")
                
                text_parts = []
                with pdfplumber.open(file_path) as pdf:
                    self.logger.info(f"PDF has {len(pdf.pages)} pages")
                    
                    for page_num, page in enumerate(pdf.pages, 1):
                        if page_num > 100:  # Limit to first 100 pages
                            break
                        
                        try:
                            page_text = page.extract_text()
                            if page_text and page_text.strip():
                                text_parts.append(page_text.strip())
                                self.logger.info(f"Page {page_num}: extracted {len(page_text)} characters")
                            else:
                                self.logger.warning(f"Page {page_num}: no text extracted")
                        except Exception as e:
                            self.logger.warning(f"Error extracting from page {page_num}: {str(e)}")
                            continue
                
                if text_parts:
                    text = "\n\n".join(text_parts)
                    self.logger.info(f"pdfplumber extracted {len(text):,} characters from {len(text_parts)} pages")
                    return text
                else:
                    self.logger.warning("pdfplumber: No text extracted from any pages")
                
            except Exception as e:
                self.logger.warning(f"pdfplumber failed: {str(e)}")
        
        # Try PyPDF2 as fallback
        if PYPDF2_AVAILABLE:
            try:
                import PyPDF2
                self.logger.info("Trying PyPDF2 for PDF extraction")
                
                text_parts = []
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    self.logger.info(f"PyPDF2: PDF has {len(pdf_reader.pages)} pages")
                    
                    for page_num, page in enumerate(pdf_reader.pages, 1):
                        if page_num > 100:  # Limit to first 100 pages
                            break
                        
                        try:
                            page_text = page.extract_text()
                            if page_text and page_text.strip():
                                text_parts.append(page_text.strip())
                                self.logger.info(f"PyPDF2 Page {page_num}: extracted {len(page_text)} characters")
                            else:
                                self.logger.warning(f"PyPDF2 Page {page_num}: no text extracted")
                        except Exception as e:
                            self.logger.warning(f"PyPDF2 error on page {page_num}: {str(e)}")
                            continue
                
                if text_parts:
                    text = "\n\n".join(text_parts)
                    self.logger.info(f"PyPDF2 extracted {len(text):,} characters from {len(text_parts)} pages")
                    return text
                else:
                    self.logger.warning("PyPDF2: No text extracted from any pages")
                
            except Exception as e:
                self.logger.warning(f"PyPDF2 failed: {str(e)}")
        
        # Try OCR as last resort if available
        try:
            import pytesseract
            from PIL import Image
            import fitz  # PyMuPDF for PDF to image conversion
            
            # Configure tesseract path for Windows
            import platform
            if platform.system() == "Windows":
                # Common Tesseract installation paths on Windows
                tesseract_paths = [
                    r"C:\Program Files\Tesseract-OCR\tesseract.exe",
                    r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
                    r"C:\Users\{}\AppData\Local\Programs\Tesseract-OCR\tesseract.exe".format(os.environ.get('USERNAME', '')),
                    r"C:\ProgramData\chocolatey\lib\tesseract\tools\tesseract.exe",
                    r"C:\ProgramData\chocolatey\bin\tesseract.exe"
                ]
                
                for path in tesseract_paths:
                    if os.path.exists(path):
                        pytesseract.pytesseract.tesseract_cmd = path
                        self.logger.info(f"Found Tesseract at: {path}")
                        break
                else:
                    self.logger.warning("Tesseract not found in common paths. Please ensure it's in your PATH.")
            
            self.logger.info("Trying OCR extraction as last resort")
            
            # Convert PDF pages to images and OCR them
            doc = fitz.open(file_path)
            text_parts = []
            
            for page_num in range(min(10, len(doc))):  # Limit to first 10 pages for OCR
                try:
                    page = doc.load_page(page_num)
                    pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))  # Higher resolution for better OCR
                    img_data = pix.tobytes("png")
                    
                    # OCR the image
                    image = Image.open(io.BytesIO(img_data))
                    
                    # Use improved OCR config for better text extraction with proper spacing
                    # PSM 6 = single uniform block, OEM 3 = default engine
                    custom_config = r'--oem 3 --psm 6'
                    
                    page_text = pytesseract.image_to_string(image, config=custom_config)
                    
                    if page_text and page_text.strip():
                        cleaned_text = page_text.strip()
                        
                        # Enhanced text cleaning and spacing correction
                        import re
                        
                        # Basic cleaning - remove excessive whitespace
                        cleaned_text = re.sub(r'\n\s*\n', '\n\n', cleaned_text)
                        cleaned_text = re.sub(r' +', ' ', cleaned_text)
                        
                        # Fix common OCR spacing issues
                        cleaned_text = self._fix_ocr_spacing(cleaned_text)
                        
                        text_parts.append(cleaned_text)
                        self.logger.info(f"OCR Page {page_num + 1}: extracted {len(page_text)} characters")
                    else:
                        self.logger.warning(f"OCR Page {page_num + 1}: no text extracted")
                        
                except Exception as e:
                    self.logger.warning(f"OCR error on page {page_num + 1}: {str(e)}")
                    continue
            
            doc.close()
            
            if text_parts:
                text = "\n\n".join(text_parts)
                self.logger.info(f"OCR extracted {len(text):,} characters from {len(text_parts)} pages")
                return text
            else:
                self.logger.warning("OCR: No text extracted from any pages")
                
        except ImportError as e:
            self.logger.warning(f"OCR libraries not available: {str(e)}")
        except Exception as e:
            self.logger.warning(f"OCR extraction failed: {str(e)}")
            
        self.logger.error("All PDF extraction methods failed")
        return None
    
    def _extract_from_docx(self, file_path: str) -> Optional[str]:
        """Extract text from modern Word document (.docx)"""
        # Try docx2txt first (usually most reliable)
        if DOCX2TXT_AVAILABLE:
            try:
                import docx2txt
                self.logger.info("Trying docx2txt for .docx extraction")
                text = docx2txt.process(file_path)
                
                if text and text.strip() and len(text.strip()) > 100:
                    self.logger.info(f"docx2txt extracted {len(text):,} characters")
                    return text.strip()
                else:
                    self.logger.warning("docx2txt extracted insufficient text")
            except Exception as e:
                self.logger.warning(f"docx2txt failed: {str(e)}")
        
        # Try python-docx as fallback
        if PYTHON_DOCX_AVAILABLE:
            try:
                from docx import Document
                self.logger.info("Trying python-docx for .docx extraction")
                
                doc = Document(file_path)
                text_parts = []
                
                # Extract paragraphs
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text_parts.append(paragraph.text.strip())
                
                # Also try to extract from tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                text_parts.append(cell.text.strip())
                
                if text_parts:
                    text = "\n\n".join(text_parts)
                    if len(text.strip()) > 100:
                        self.logger.info(f"python-docx extracted {len(text):,} characters")
                        return text
                    else:
                        self.logger.warning("python-docx extracted insufficient text")
                
            except Exception as e:
                self.logger.warning(f"python-docx failed: {str(e)}")
        
        # If both methods failed, try treating as ZIP and extracting XML content
        try:
            self.logger.info("Trying to extract content directly from .docx ZIP structure")
            import zipfile
            import xml.etree.ElementTree as ET
            
            with zipfile.ZipFile(file_path, 'r') as docx_zip:
                # Try to read document.xml
                try:
                    xml_content = docx_zip.read('word/document.xml')
                    root = ET.fromstring(xml_content)
                    
                    # Extract text from XML elements
                    text_parts = []
                    for elem in root.iter():
                        if elem.text and elem.text.strip():
                            text_parts.append(elem.text.strip())
                    
                    if text_parts:
                        text = " ".join(text_parts)
                        if len(text.strip()) > 100:
                            self.logger.info(f"Direct XML extraction got {len(text):,} characters")
                            return text
                        
                except Exception as e:
                    self.logger.warning(f"Direct XML extraction failed: {str(e)}")
        
        except Exception as e:
            self.logger.warning(f"ZIP-based extraction failed: {str(e)}")
        
        # Last resort - try raw text extraction
        self.logger.warning("All .docx methods failed, trying raw text extraction")
        return self._extract_as_text(file_path)
    
    def _extract_from_old_doc(self, file_path: str) -> Optional[str]:
        """Extract text from old Word document (.doc)"""
        # Try antiword command (if available)
        try:
            self.logger.info("Trying antiword command for .doc extraction")
            import subprocess
            result = subprocess.run(
                ['antiword', file_path], 
                capture_output=True, 
                text=True, 
                timeout=30
            )
            
            if result.returncode == 0 and result.stdout.strip():
                text = result.stdout.strip()
                self.logger.info(f"antiword extracted {len(text):,} characters")
                return text
            
        except (FileNotFoundError, subprocess.TimeoutExpired, Exception) as e:
            self.logger.warning(f"antiword failed: {str(e)}")
        
        # Try reading as binary and looking for embedded content
        return self._extract_embedded_content(file_path)
    
    def _extract_embedded_pdf(self, file_path: str) -> Optional[str]:
        """Extract PDF content embedded in another document"""
        try:
            self.logger.info("Attempting to extract embedded PDF content")
            
            with open(file_path, 'rb') as f:
                content = f.read()
            
            # Find PDF start and end markers
            pdf_start = content.find(b'%PDF')
            if pdf_start == -1:
                self.logger.warning("No PDF start marker found")
                return None
            
            # Look for PDF end marker or use end of file
            pdf_end_markers = [b'%%EOF', b'\nendstream\n', b'\r\nendstream\r\n']
            pdf_end = len(content)
            
            for marker in pdf_end_markers:
                marker_pos = content.rfind(marker)
                if marker_pos > pdf_start:
                    pdf_end = marker_pos + len(marker)
                    break
            
            # Extract PDF content
            pdf_content = content[pdf_start:pdf_end]
            
            # Write to temporary PDF file
            import tempfile
            with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as temp_pdf:
                temp_pdf.write(pdf_content)
                temp_pdf_path = temp_pdf.name
            
            try:
                # Extract text from the temporary PDF
                text = self._extract_from_pdf(temp_pdf_path)
                return text
            finally:
                # Clean up temporary file
                try:
                    os.unlink(temp_pdf_path)
                except:
                    pass
            
        except Exception as e:
            self.logger.error(f"Error extracting embedded PDF: {str(e)}")
            return None
    
    def _extract_embedded_content(self, file_path: str) -> Optional[str]:
        """Extract content from compound documents that might contain embedded files"""
        try:
            self.logger.info("Attempting to extract embedded content from compound document")
            
            with open(file_path, 'rb') as f:
                content = f.read()
            
            # Check if this contains an embedded PDF
            if b'%PDF' in content:
                return self._extract_embedded_pdf(file_path)
            
            # Otherwise try raw text extraction with better filtering
            return self._extract_as_text(file_path)
            
        except Exception as e:
            self.logger.warning(f"Error extracting embedded content: {str(e)}")
            return None
    
    def _extract_as_text(self, file_path: str) -> Optional[str]:
        """Extract readable text from file using raw text reading"""
        try:
            self.logger.info("Using raw text reading method")
            
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding, errors='ignore') as f:
                        content = f.read()
                    
                    # Filter for printable characters and common patterns
                    import string
                    
                    # Keep letters, digits, punctuation, and whitespace
                    printable_chars = string.ascii_letters + string.digits + string.punctuation + string.whitespace
                    filtered_content = ''.join(char for char in content if char in printable_chars)
                    
                    # Look for text patterns that suggest legal document content
                    legal_patterns = [
                        'tribunal', 'appeal', 'decision', 'judgment', 'court',
                        'paragraph', 'applicant', 'respondent', 'immigration',
                        'asylum', 'secretary of state', 'home office'
                    ]
                    
                    # Count how many legal terms we find
                    content_lower = filtered_content.lower()
                    legal_term_count = sum(1 for pattern in legal_patterns if pattern in content_lower)
                    
                    # If we found some legal terms and have reasonable content length
                    if legal_term_count >= 2 and len(filtered_content.strip()) > 500:
                        # Clean up the text further
                        lines = filtered_content.split('\n')
                        cleaned_lines = []
                        
                        for line in lines:
                            line = line.strip()
                            # Keep lines that have reasonable length and legal-sounding content
                            if (len(line) > 10 and 
                                any(term in line.lower() for term in legal_patterns + ['section', 'rule', 'act', 'law'])):
                                cleaned_lines.append(line)
                        
                        if cleaned_lines:
                            text = '\n'.join(cleaned_lines)
                            self.logger.info(f"Raw text extraction with {encoding} extracted {len(text):,} characters")
                            return text
                        
                except UnicodeDecodeError:
                    continue
            
            self.logger.warning("Raw text extraction did not find sufficient legal content")
            return None
            
        except Exception as e:
            self.logger.error(f"Error in raw text extraction: {str(e)}")
            return None
    
    def extract_text_from_file(self, file_path: Path) -> Optional[str]:
        """Extract text from downloaded file based on extension"""
        if not file_path.exists():
            return None
        
        # Use the new unified extraction method
        return self.extract_text_from_document(str(file_path), str(file_path))
    
    def get_json_filename(self, reference_number: str) -> str:
        """Convert reference number to JSON filename"""
        # Replace forward slashes with underscores to match existing naming convention
        safe_ref = reference_number.replace('/', '_')
        return f"{safe_ref}.json"
    
    def update_json_file(self, reference_number: str, decision_text: str) -> bool:
        """
        Update the corresponding JSON file with extracted decision text
        
        Args:
            reference_number: Case reference number
            decision_text: Extracted decision text
            
        Returns:
            True if successful, False otherwise
        """
        try:
            json_filename = self.get_json_filename(reference_number)
            json_path = self.json_dir / json_filename
            
            if not json_path.exists():
                self.logger.error(f"JSON file not found: {json_path}")
                return False
            
            # Load existing JSON
            with open(json_path, 'r', encoding='utf-8') as f:
                case_data = json.load(f)
            
            # Update with extracted text
            case_data['decision_text'] = decision_text
            case_data['text_extraction_date'] = datetime.now().isoformat()
            case_data['text_extraction_method'] = 'document_download'
            
            # Save updated JSON
            with open(json_path, 'w', encoding='utf-8') as f:
                json.dump(case_data, f, indent=2, ensure_ascii=False)
            
            self.logger.info(f"Updated JSON file: {json_filename}")
            return True
            
        except Exception as e:
            self.logger.error(f"Error updating JSON for {reference_number}: {str(e)}")
            return False
    
    def process_single_case(self, case_row: pd.Series) -> Tuple[bool, str]:
        """
        Process a single case: get fresh URLs, download document and extract text
        
        Args:
            case_row: Pandas Series containing case data
            
        Returns:
            Tuple of (success, status_message)
        """
        reference_number = case_row['reference_number']
        case_url = case_row.get('url')
        
        self.logger.info(f"Processing case: {reference_number}")
        
        if not case_url or pd.isna(case_url):
            return False, "No main case URL available"
        
        # Get fresh download URLs
        word_url, pdf_url = self.get_fresh_download_urls(case_url)
        
        if not word_url and not pdf_url:
            return False, "No download URLs found on case page"
        
        # Track extraction attempts
        extraction_attempts = []
        
        # Try Word document first, then PDF
        for url, file_type in [(word_url, 'word'), (pdf_url, 'pdf')]:
            if not url:
                continue
            
            # Determine file extension from URL or file type
            if file_type == 'word':
                extension = '.doc'
            else:
                extension = '.pdf'
            
            # Create filename
            safe_ref = reference_number.replace('/', '_').replace('[', '').replace(']', '')
            filename = f"{safe_ref}_{file_type}{extension}"
            
            # Download file
            downloaded_file = self.download_file(url, filename)
            if not downloaded_file:
                extraction_attempts.append(f"{file_type}: download failed")
                continue
            
            # Extract text
            self.logger.info(f"Attempting text extraction from {file_type} file")
            extracted_text = self.extract_text_from_file(downloaded_file)
            
            # Clean up downloaded file
            try:
                downloaded_file.unlink()
            except:
                pass
            
            if extracted_text and len(extracted_text.strip()) > 100:  # Minimum viable text length
                # Update JSON file
                if self.update_json_file(reference_number, extracted_text):
                    extraction_method = self._determine_extraction_method(str(downloaded_file), extracted_text)
                    return True, f"Successfully extracted {len(extracted_text):,} characters from {file_type} using {extraction_method}"
                else:
                    return False, f"Text extracted but failed to update JSON"
            else:
                extraction_attempts.append(f"{file_type}: extraction failed (text too short)")
                self.logger.warning(f"Extracted text too short or empty from {file_type} for {reference_number}")
        
        # If we get here, all attempts failed
        attempts_summary = "; ".join(extraction_attempts)
        return False, f"All extraction attempts failed: {attempts_summary}"
    
    def _determine_extraction_method(self, file_path: str, extracted_text: str) -> str:
        """
        Determine which extraction method was likely used based on content characteristics
        """
        # Look for OCR characteristics (common OCR artifacts)
        ocr_indicators = [
            'l ' in extracted_text[:1000],  # OCR often confuses 'I' with 'l '
            ' rn ' in extracted_text[:1000],  # OCR often confuses 'm' with 'rn'
            extracted_text.count('\n') > len(extracted_text) // 50,  # OCR tends to have many line breaks
        ]
        
        if sum(ocr_indicators) >= 2:
            return "OCR"
        elif file_path.endswith('.pdf'):
            return "PDF text extraction"
        elif any(ext in file_path.lower() for ext in ['.doc', '.docx']):
            return "Word document extraction"
        else:
            return "Raw text extraction"
    
    def process_all_cases(self) -> Dict[str, int]:
        """
        Process all cases with missing decision text
        
        Returns:
            Dictionary with processing statistics
        """
        if len(self.missing_cases) == 0:
            self.logger.warning("No cases to process!")
            return {'total': 0, 'successful': 0, 'failed': 0}
        
        stats = {
            'total': len(self.missing_cases),
            'successful': 0,
            'failed': 0,
            'start_time': datetime.now().isoformat()
        }
        
        self.logger.info(f"Starting to process {stats['total']} cases")
        
        try:
            for idx, case_row in self.missing_cases.iterrows():
                reference_number = case_row['reference_number']
                
                try:
                    success, message = self.process_single_case(case_row)
                    
                    if success:
                        stats['successful'] += 1
                        self.logger.info(f"SUCCESS {reference_number}: {message}")
                    else:
                        stats['failed'] += 1
                        self.logger.warning(f"FAILED {reference_number}: {message}")
                    
                    # Log progress every 10 cases
                    processed = stats['successful'] + stats['failed']
                    if processed % 10 == 0:
                        self.logger.info(f"Progress: {processed}/{stats['total']} processed "
                                       f"({stats['successful']} successful, {stats['failed']} failed)")
                
                except Exception as e:
                    stats['failed'] += 1
                    self.logger.error(f"Error processing {reference_number}: {str(e)}")
        
        except KeyboardInterrupt:
            self.logger.info("Processing interrupted by user")
        
        finally:
            stats['end_time'] = datetime.now().isoformat()
            
            # Cleanup temp directory
            try:
                import shutil
                shutil.rmtree(self.temp_dir)
                self.logger.info(f"Cleaned up temporary directory: {self.temp_dir}")
            except:
                self.logger.warning(f"Could not clean up temporary directory: {self.temp_dir}")
        
        return stats
    
    def test_single_download(self, reference_number: str = None) -> bool:
        """
        Test processing a single case for verification
        
        Args:
            reference_number: Specific case to test, or None for first available
            
        Returns:
            True if successful, False otherwise
        """
        if len(self.missing_cases) == 0:
            self.logger.error("No cases available for testing")
            return False
        
        if reference_number:
            test_case = self.missing_cases[self.missing_cases['reference_number'] == reference_number]
            if len(test_case) == 0:
                self.logger.error(f"Case {reference_number} not found in missing cases")
                return False
            test_case = test_case.iloc[0]
        else:
            test_case = self.missing_cases.iloc[0]
        
        self.logger.info(f"Testing single download for: {test_case['reference_number']}")
        
        success, message = self.process_single_case(test_case)
        
        if success:
            self.logger.info(f"SUCCESS: Test successful: {message}")
        else:
            self.logger.error(f"FAILED: Test failed: {message}")
        
        return success

    def _fix_ocr_spacing(self, text: str) -> str:
        """
        Fix common OCR spacing issues where words are concatenated without spaces
        """
        import re
        
        # Pattern to add spaces before capital letters that follow lowercase letters
        # e.g., "HeardatField" -> "Heard at Field"
        text = re.sub(r'([a-z])([A-Z])', r'\1 \2', text)
        
        # Pattern to add spaces before numbers that follow letters
        # e.g., "Appeal No:HX471962001" -> "Appeal No: HX47196 2001"
        text = re.sub(r'([a-zA-Z])(\d)', r'\1 \2', text)
        
        # Pattern to add spaces after numbers that are followed by letters (except common patterns)
        # e.g., "2001On" -> "2001 On", but keep "2002]" as is
        text = re.sub(r'(\d)([A-Z][a-z])', r'\1 \2', text)
        
        # Add spaces around common legal terms that often get concatenated
        legal_patterns = [
            (r'([a-z])(APPELLANT|RESPONDENT|DETERMINATION|REASONS)', r'\1 \2'),
            (r'(APPELLANT|RESPONDENT|DETERMINATION|REASONS)([a-z])', r'\1 \2'),
            (r'([a-z])(Secretary|Department|Court|Tribunal|Appeal)', r'\1 \2'),
            (r'(Secretary|Department|Court|Tribunal|Appeal)([A-Z][a-z])', r'\1 \2'),
            (r'([a-z])(Home|Office|State|Kingdom)', r'\1 \2'),
            (r'(Home|Office|State|Kingdom)([A-Z][a-z])', r'\1 \2'),
        ]
        
        for pattern, replacement in legal_patterns:
            text = re.sub(pattern, replacement, text)
        
        # Fix common date/time patterns
        # e.g., "8October2002" -> "8 October 2002"
        months = ['January', 'February', 'March', 'April', 'May', 'June',
                 'July', 'August', 'September', 'October', 'November', 'December']
        
        for month in months:
            text = re.sub(f'(\\d)({month})', r'\1 \2', text)
            text = re.sub(f'({month})(\\d)', r'\1 \2', text)
        
        # Add spaces around punctuation where missing
        text = re.sub(r'([a-zA-Z])([.,:;!?])([A-Z])', r'\1\2 \3', text)
        
        # Fix common title patterns
        # e.g., "MrMohammed" -> "Mr Mohammed", "MrsJAJC" -> "Mrs J A J C"
        titles = ['Mr', 'Mrs', 'Ms', 'Dr', 'Prof', 'Hon', 'Rt']
        for title in titles:
            text = re.sub(f'({title})([A-Z][a-z])', r'\1 \2', text)
            text = re.sub(f'([a-z])({title})', r'\1 \2', text)
        
        # Clean up any double spaces that might have been introduced
        text = re.sub(r' +', ' ', text)
        
        return text

    def save_checkpoint(self):
        """Save current progress to checkpoint file"""
        checkpoint_data = {
            'timestamp': datetime.now().isoformat(),
            'current_batch': self.current_batch,
            'processed_cases': list(self.processed_cases),
            'successful_cases': self.successful_cases.copy(),
            'failed_cases': self.failed_cases.copy(),
            'cases_processed': self.cases_processed,
            'total_download_time': self.total_download_time,
            'total_extraction_time': self.total_extraction_time,
            'current_rate_limit': self.rate_limit
        }
        
        checkpoint_file = self.checkpoint_dir / f"progress_checkpoint.json"
        with open(checkpoint_file, 'w', encoding='utf-8') as f:
            json.dump(checkpoint_data, f, indent=2)
        
        self.logger.info(f"💾 Checkpoint saved: {len(self.processed_cases)} cases processed")
    
    def load_checkpoint(self) -> bool:
        """Load progress from checkpoint file"""
        checkpoint_file = self.checkpoint_dir / "progress_checkpoint.json"
        
        if not checkpoint_file.exists():
            self.logger.info("🆕 No checkpoint found - starting fresh")
            return False
        
        try:
            with open(checkpoint_file, 'r', encoding='utf-8') as f:
                checkpoint_data = json.load(f)
            
            self.current_batch = checkpoint_data.get('current_batch', 0)
            self.processed_cases = set(checkpoint_data.get('processed_cases', []))
            self.successful_cases = checkpoint_data.get('successful_cases', [])
            self.failed_cases = checkpoint_data.get('failed_cases', [])
            self.cases_processed = checkpoint_data.get('cases_processed', 0)
            self.total_download_time = checkpoint_data.get('total_download_time', 0)
            self.total_extraction_time = checkpoint_data.get('total_extraction_time', 0)
            self.rate_limit = checkpoint_data.get('current_rate_limit', self.initial_rate_limit)
            
            self.logger.info(f"📂 Checkpoint loaded: {len(self.processed_cases)} cases already processed")
            self.logger.info(f"🔄 Resuming from batch {self.current_batch}")
            return True
            
        except Exception as e:
            self.logger.error(f"❌ Error loading checkpoint: {str(e)}")
            return False
    
    def adjust_rate_limit(self, success: bool, response_time: float):
        """Adaptively adjust rate limiting based on performance - less aggressive"""
        if success and response_time < 1.0:  # More lenient success threshold
            # Speed up if responses are fast
            self.rate_limit = max(0.2, self.rate_limit * 0.98)  # Slower reduction
        elif not success or response_time > 5.0:  # Higher threshold for slowdown
            # Slow down if getting errors or very slow responses
            self.rate_limit = min(2.0, self.rate_limit * 1.1)  # Smaller increase, lower cap
        
        # Log significant changes
        if abs(self.rate_limit - self.initial_rate_limit) > 0.3:
            self.logger.info(f"⚡ Rate limit adjusted to {self.rate_limit:.2f}s")
    
    def cleanup_memory(self):
        """Clean up memory after heavy operations"""
        import gc
        gc.collect()
        
        # Log memory usage if psutil is available
        try:
            import psutil
            process = psutil.Process()
            memory_mb = process.memory_info().rss / 1024 / 1024
            if memory_mb > 1000:  # Log if using over 1GB
                self.logger.info(f"🧠 Memory usage: {memory_mb:.1f} MB")
        except ImportError:
            pass

    def process_all_cases_optimized(self, resume: bool = True) -> Dict[str, int]:
        """
        Process all cases with optimized batch processing and checkpointing
        
        Args:
            resume: Whether to resume from checkpoint if available
            
        Returns:
            Dictionary with processing statistics
        """
        self.start_time = datetime.now()
        
        # Load checkpoint if resuming
        if resume:
            self.load_checkpoint()
        
        if len(self.missing_cases) == 0:
            self.logger.warning("No cases to process!")
            return {'total': 0, 'successful': 0, 'failed': 0}
        
        # Filter out already processed cases if resuming
        remaining_cases = self.missing_cases.copy()
        if self.processed_cases:
            mask = ~remaining_cases['reference_number'].isin(self.processed_cases)
            remaining_cases = remaining_cases[mask]
            self.logger.info(f"📋 {len(self.processed_cases)} cases already processed, {len(remaining_cases)} remaining")
        
        total_cases = len(remaining_cases)
        
        # Debug logging for key variables
        self.logger.info(f"🔍 DEBUG: total_missing_cases = {len(self.missing_cases)}")
        self.logger.info(f"🔍 DEBUG: already_processed = {len(self.processed_cases)}")
        self.logger.info(f"🔍 DEBUG: remaining_cases = {total_cases}")
        self.logger.info(f"🔍 DEBUG: batch_size = {self.batch_size}")
        
        if total_cases == 0:
            self.logger.info("✅ All cases already processed!")
            return {
                'total': len(self.missing_cases),
                'successful': len(self.successful_cases),
                'failed': len(self.failed_cases),
                'successful_cases': self.successful_cases
            }
        
        self.logger.info(f"🚀 Starting optimized processing of {total_cases} cases in batches of {self.batch_size}")
        
        # Console notification about logging
        print(f"\n🚀 Starting optimized processing of {total_cases} cases")
        print(f"📊 Batch size: {self.batch_size}")
        print(f"🔇 Console logging: QUIET mode (only key progress + warnings/errors)")
        print(f"📄 Detailed logs: logs/document_extractor_*.log")
        print(f"{'='*60}")
        
        # Add debug logging for batch calculation
        total_batches = (total_cases + self.batch_size - 1) // self.batch_size
        self.logger.info(f"🔢 Calculated {total_batches} batches for {total_cases} cases with batch size {self.batch_size}")
        
        # Debug: Show the actual range values
        batch_ranges = list(range(0, total_cases, self.batch_size))
        self.logger.info(f"🔍 DEBUG: batch_ranges = {batch_ranges[:10]}{'...' if len(batch_ranges) > 10 else ''} (showing first 10)")
        self.logger.info(f"🔍 DEBUG: batch_ranges length = {len(batch_ranges)}")
        
        # Create overall progress bar
        overall_pbar = tqdm(
            total=total_cases,
            desc="Processing cases",
            unit="cases",
            position=0,
            leave=True,
            bar_format="{desc}: {percentage:3.0f}%|{bar}| {n_fmt}/{total_fmt} [{elapsed}<{remaining}, {rate_fmt}]"
        )
        
        try:
            # Process in batches - add debug logging
            batch_count = 0
            self.logger.info(f"🔍 DEBUG: Starting for loop with range(0, {total_cases}, {self.batch_size})")
            for batch_start in range(0, total_cases, self.batch_size):
                batch_count += 1
                batch_end = min(batch_start + self.batch_size, total_cases)
                
                # Debug: Log the actual slice being taken
                self.logger.info(f"🔍 DEBUG: Slicing remaining_cases.iloc[{batch_start}:{batch_end}]")
                batch_cases = remaining_cases.iloc[batch_start:batch_end]
                
                self.current_batch += 1
                batch_size_actual = len(batch_cases)
                
                self.logger.info(f"🔄 BATCH LOOP: Starting batch {batch_count}/{total_batches}")
                self.logger.info(f"📦 Processing batch {self.current_batch}: cases {batch_start + 1}-{batch_end} ({batch_size_actual} cases)")
                
                # Validate batch processing parameters
                if batch_size_actual <= 0:
                    self.logger.error(f"❌ Invalid batch size: {batch_size_actual}")
                    continue
                
                if batch_start >= total_cases:
                    self.logger.error(f"❌ Batch start {batch_start} exceeds total cases {total_cases}")
                    break
                
                batch_start_time = time.time()
                batch_successful = 0
                batch_failed = 0
                
                # Create batch progress bar
                batch_pbar = tqdm(
                    total=batch_size_actual,
                    desc=f"Batch {self.current_batch}",
                    unit="cases",
                    position=1,
                    leave=False,
                    bar_format="{desc}: {percentage:3.0f}%|{bar}| {n_fmt}/{total_fmt} [{elapsed}, {rate_fmt}]"
                )
                
                # Process each case in the batch
                batch_error_count = 0
                for idx, case_row in batch_cases.iterrows():
                    reference_number = case_row['reference_number']
                    
                    # Skip if already processed (safety check)
                    if reference_number in self.processed_cases:
                        batch_pbar.update(1)
                        overall_pbar.update(1)
                        continue
                    
                    try:
                        case_start_time = time.time()
                        success, message = self.process_single_case_optimized(case_row)
                        case_time = time.time() - case_start_time
                        
                        # Track performance
                        self.cases_processed += 1
                        self.processed_cases.add(reference_number)
                        
                        if success:
                            batch_successful += 1
                            self.successful_cases.append(reference_number)
                            self.logger.info(f"✅ {reference_number}: {message} ({case_time:.1f}s)")
                            batch_pbar.set_postfix_str(f"✅ {reference_number}")
                        else:
                            batch_failed += 1
                            self.failed_cases.append(reference_number)
                            self.logger.warning(f"❌ {reference_number}: {message}")
                            batch_pbar.set_postfix_str(f"❌ {reference_number}")
                        
                        # Update progress bars
                        batch_pbar.update(1)
                        overall_pbar.update(1)
                        
                        # Update overall progress bar description with statistics
                        success_rate = len(self.successful_cases) / len(self.processed_cases) * 100 if self.processed_cases else 0
                        overall_pbar.set_description(f"Processing cases (Success: {success_rate:.1f}%)")
                        
                        # Memory cleanup every 50 cases
                        if self.cases_processed % 50 == 0:
                            self.cleanup_memory()
                    
                    except Exception as e:
                        batch_failed += 1
                        batch_error_count += 1
                        self.failed_cases.append(reference_number)
                        self.processed_cases.add(reference_number)
                        self.logger.error(f"💥 {reference_number}: Unexpected error: {str(e)}")
                        
                        # Update progress bars
                        batch_pbar.set_postfix_str(f"💥 {reference_number}")
                        batch_pbar.update(1)
                        overall_pbar.update(1)
                        
                        # If too many errors in this batch, consider stopping
                        if batch_error_count > batch_size_actual * 0.8:  # More than 80% errors
                            self.logger.warning(f"⚠️  High error rate in batch {self.current_batch} ({batch_error_count}/{batch_size_actual})")
                
                # Close batch progress bar
                batch_pbar.close()
                
                # Batch completion
                batch_time = time.time() - batch_start_time
                batch_rate = batch_size_actual / batch_time if batch_time > 0 else 0
                
                # Console progress update (using print to bypass reduced console logging)
                print(f"\n📊 Batch {self.current_batch}/{total_batches} complete: {batch_successful} successful, {batch_failed} failed ({batch_rate:.2f} cases/min)")
                self.logger.info(f"📊 Batch {self.current_batch} complete: {batch_successful} successful, {batch_failed} failed ({batch_rate:.2f} cases/min)")
                
                # Save checkpoint after each batch
                try:
                    self.save_checkpoint()
                except Exception as e:
                    self.logger.error(f"❌ Failed to save checkpoint: {str(e)}")
                
                # Progress report
                total_processed = len(self.processed_cases)
                overall_progress = (total_processed / len(self.missing_cases)) * 100
                
                # Console progress report  
                print(f"📈 Overall progress: {total_processed}/{len(self.missing_cases)} cases ({overall_progress:.1f}%)")
                self.logger.info(f"📈 Overall progress: {total_processed}/{len(self.missing_cases)} cases ({overall_progress:.1f}%)")
                
                # Debug: Log loop continuation
                remaining_batches = total_batches - batch_count
                self.logger.info(f"🔄 BATCH LOOP: Completed batch {batch_count}/{total_batches}, {remaining_batches} batches remaining")
                
                # Estimate time remaining
                if self.start_time and total_processed > 0:
                    elapsed = datetime.now() - self.start_time
                    rate = total_processed / elapsed.total_seconds()
                    remaining_cases_calc = len(self.missing_cases) - total_processed
                    eta_seconds = remaining_cases_calc / rate if rate > 0 else 0
                    eta = datetime.now() + timedelta(seconds=eta_seconds)
                    
                    # Console ETA update
                    print(f"⏱️  ETA: {eta.strftime('%H:%M:%S')} ({eta_seconds/3600:.1f} hours remaining)")
                    self.logger.info(f"⏱️  ETA: {eta.strftime('%Y-%m-%d %H:%M:%S')} ({eta_seconds/3600:.1f} hours remaining)")
                    
                    # Update overall progress bar with ETA
                    eta_str = f"ETA: {eta.strftime('%H:%M:%S')}"
                    overall_pbar.set_postfix_str(eta_str)
                
                # Verify we haven't somehow processed all cases
                if total_processed >= len(self.missing_cases):
                    print(f"🎉 All cases processed! Breaking from batch loop.")
                    self.logger.info(f"🎉 All cases processed! Breaking from batch loop.")
                    break
            
            # Debug: Log when loop completes
            self.logger.info(f"🏁 BATCH LOOP: Completed {batch_count} out of {total_batches} batches!")
        
        except KeyboardInterrupt:
            self.logger.info("⚠️  Processing interrupted by user - progress saved")
        
        finally:
            # Close progress bar
            overall_pbar.close()
            
            # Final statistics
            final_stats = {
                'total': len(self.missing_cases),
                'successful': len(self.successful_cases),
                'failed': len(self.failed_cases),
                'start_time': self.start_time.isoformat() if self.start_time else None,
                'end_time': datetime.now().isoformat(),
                'successful_cases': self.successful_cases
            }
            
            # Save final checkpoint
            self.save_checkpoint()
            
            # Cleanup temp directory
            try:
                import shutil
                shutil.rmtree(self.temp_dir)
                self.logger.info(f"🧹 Cleaned up temporary directory: {self.temp_dir}")
            except:
                self.logger.warning(f"⚠️  Could not clean up temporary directory: {self.temp_dir}")
            
            return final_stats

    def process_single_case_optimized(self, case_row: pd.Series) -> Tuple[bool, str]:
        """
        Process a single case with performance tracking and adaptive rate limiting
        
        Args:
            case_row: Pandas Series containing case data
            
        Returns:
            Tuple of (success, status_message)
        """
        reference_number = case_row['reference_number']
        case_url = case_row.get('url')
        
        if not case_url or pd.isna(case_url):
            return False, "No main case URL available"
        
        # Performance tracking
        download_start = time.time()
        
        # Get fresh download URLs
        word_url, pdf_url = self.get_fresh_download_urls(case_url)
        
        if not word_url and not pdf_url:
            return False, "No download URLs found on case page"
        
        # Track extraction attempts
        extraction_attempts = []
        
        # Try Word document first, then PDF
        for url, file_type in [(word_url, 'word'), (pdf_url, 'pdf')]:
            if not url:
                continue
            
            # Determine file extension from URL or file type
            if file_type == 'word':
                extension = '.doc'
            else:
                extension = '.pdf'
            
            # Create filename
            safe_ref = reference_number.replace('/', '_').replace('[', '').replace(']', '')
            filename = f"{safe_ref}_{file_type}{extension}"
            
            # Download file with performance tracking
            download_start_time = time.time()
            downloaded_file = self.download_file_optimized(url, filename)
            download_time = time.time() - download_start_time
            
            # Adjust rate limiting based on download performance
            if downloaded_file:
                self.adjust_rate_limit(True, download_time)
                self.total_download_time += download_time
            else:
                self.adjust_rate_limit(False, download_time)
                extraction_attempts.append(f"{file_type}: download failed")
                continue
            
            # Extract text with performance tracking
            extraction_start_time = time.time()
            extracted_text = self.extract_text_from_file(downloaded_file)
            extraction_time = time.time() - extraction_start_time
            self.total_extraction_time += extraction_time
            
            # Clean up downloaded file
            try:
                downloaded_file.unlink()
            except:
                pass
            
            if extracted_text and len(extracted_text.strip()) > 100:  # Minimum viable text length
                # Update JSON file
                if self.update_json_file(reference_number, extracted_text):
                    extraction_method = self._determine_extraction_method(str(downloaded_file), extracted_text)
                    total_time = time.time() - download_start
                    return True, f"Successfully extracted {len(extracted_text):,} characters from {file_type} using {extraction_method} ({total_time:.1f}s)"
                else:
                    return False, f"Text extracted but failed to update JSON"
            else:
                extraction_attempts.append(f"{file_type}: extraction failed (text too short)")
        
        # If we get here, all attempts failed
        attempts_summary = "; ".join(extraction_attempts)
        return False, f"All extraction attempts failed: {attempts_summary}"
    
    def download_file_optimized(self, url: str, filename: str) -> Optional[Path]:
        """
        Optimized download with adaptive rate limiting and better error handling
        
        Args:
            url: Download URL
            filename: Local filename to save as
            
        Returns:
            Path to downloaded file or None if failed
        """
        if not url or pd.isna(url):
            return None
        
        file_path = self.temp_dir / filename
        
        for attempt in range(self.max_retries + 1):
            try:
                # Adaptive jitter based on current rate limit
                jitter = random.uniform(-0.1, 0.1) * self.rate_limit
                actual_delay = max(0.1, self.rate_limit + jitter)
                time.sleep(actual_delay)
                
                request_start = time.time()
                response = self.session.get(url, timeout=60, stream=True)
                request_time = time.time() - request_start
                
                # Handle rate limiting
                if response.status_code == 429:
                    wait_time = int(response.headers.get('Retry-After', 60))
                    self.logger.warning(f"⏸️  Rate limited. Waiting {wait_time} seconds.")
                    time.sleep(wait_time)
                    # Increase rate limit for future requests
                    self.rate_limit = min(5.0, self.rate_limit * 1.5)
                    continue
                
                response.raise_for_status()
                
                # Check if we got HTML instead of a document (redirect to login page, etc.)
                content_type = response.headers.get('content-type', '').lower()
                if 'text/html' in content_type:
                    return None
                
                # Save file
                with open(file_path, 'wb') as f:
                    for chunk in response.iter_content(chunk_size=8192):
                        if chunk:
                            f.write(chunk)
                
                file_size = file_path.stat().st_size
                
                # Basic file validation
                if file_size < 100:  # Suspiciously small file
                    return None
                
                # Adjust rate limit based on successful request time
                self.adjust_rate_limit(True, request_time)
                
                return file_path
                
            except Exception as e:
                if attempt < self.max_retries:
                    wait_time = (2 ** attempt) + random.uniform(0, 1)
                    time.sleep(wait_time)
                else:
                    self.adjust_rate_limit(False, 5.0)  # Penalize rate limit for failures
                    return None

if __name__ == "__main__":
    # Example usage
    extractor = DocumentTextExtractor(
        csv_path="preprocessing/processed_data/processed_legal_cases.csv",
        test_mode=True,
        test_count=5
    )
    
    # Test single download
    extractor.test_single_download() 